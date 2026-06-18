"""Google Drive integration.

Pulls documents from a shared team Drive folder (read-only via service account)
and extracts plain text suitable for chunking + embedding.

Three supported source types:
- Google Docs (`application/vnd.google-apps.document`): exported as text/plain
- PDFs (`application/pdf`): extracted via pypdf
- Word (`application/vnd.openxmlformats-officedocument.wordprocessingml.document`): extracted via python-docx

Extractors are pure functions (bytes -> text) so they can be unit-tested without
hitting Drive. Drive I/O takes the Drive service as a parameter so tests can
inject a fake.

Configure via `[google].service_account_path` and `[google].shared_drive_folder_id`
in `.streamlit/secrets.toml`. Service account must have read access to the folder.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass, field
from datetime import datetime
from functools import lru_cache
from pathlib import Path
from typing import Any, Iterable

ROOT = Path(__file__).resolve().parent.parent

GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
GOOGLE_FOLDER_MIME = "application/vnd.google-apps.folder"
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

SUPPORTED_MIMES = {GOOGLE_DOC_MIME, PDF_MIME, DOCX_MIME}

# Fields requested from Drive's files.list/get for change detection + dispatch.
_FILE_FIELDS = "id,name,mimeType,modifiedTime,parents,trashed"


@dataclass
class DriveDocMeta:
    file_id: str
    name: str
    mime_type: str
    modified_time: datetime
    parents: list[str] = field(default_factory=list)


@dataclass
class DriveDoc:
    file_id: str
    name: str
    mime_type: str
    modified_time: datetime
    parents: list[str]
    text: str

    @property
    def meta(self) -> DriveDocMeta:
        return DriveDocMeta(
            file_id=self.file_id,
            name=self.name,
            mime_type=self.mime_type,
            modified_time=self.modified_time,
            parents=list(self.parents),
        )


# ----------------------- Configuration -----------------------


def _drive_secrets() -> dict:
    try:
        import streamlit as st
        return dict(st.secrets["google"])
    except Exception:
        pass
    secrets_path = ROOT / ".streamlit" / "secrets.toml"
    if secrets_path.exists():
        import toml
        return toml.load(secrets_path).get("google", {})
    return {}


def _has_credentials(s: dict) -> bool:
    """True if any supported credential form is present.

    Three forms, in priority order:
    - `service_account_info`: a TOML table (Streamlit Cloud-friendly: paste the
      JSON's fields as a nested `[google.service_account_info]` section).
    - `service_account_json`: the raw JSON key as a single string.
    - `service_account_path`: a path to a JSON key file (local dev).
    """
    return bool(
        s.get("service_account_info")
        or s.get("service_account_json")
        or s.get("service_account_path")
    )


def _is_configured() -> bool:
    s = _drive_secrets()
    return _has_credentials(s) and bool(s.get("shared_drive_folder_id"))


# ----------------------- Drive service factory -----------------------


_SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]


def _load_credentials(s: dict) -> Any:
    """Build service-account Credentials from whichever form is configured."""
    from google.oauth2 import service_account

    info = s.get("service_account_info")
    if info:
        return service_account.Credentials.from_service_account_info(
            dict(info), scopes=_SCOPES
        )

    raw = s.get("service_account_json")
    if raw:
        import json
        return service_account.Credentials.from_service_account_info(
            json.loads(raw), scopes=_SCOPES
        )

    sa_path = s["service_account_path"]
    if not Path(sa_path).is_absolute():
        sa_path = ROOT / sa_path
    return service_account.Credentials.from_service_account_file(str(sa_path), scopes=_SCOPES)


@lru_cache(maxsize=1)
def _drive_service() -> Any:
    """Build a Drive v3 service from the configured service account credentials.

    Raises RuntimeError if Drive isn't configured (use `_is_configured()` to gate).
    """
    if not _is_configured():
        raise RuntimeError(
            "Google Drive isn't configured. Set a [google] credential "
            "(service_account_info / service_account_json / service_account_path) "
            "and shared_drive_folder_id in .streamlit/secrets.toml."
        )
    from googleapiclient.discovery import build

    creds = _load_credentials(_drive_secrets())
    return build("drive", "v3", credentials=creds, cache_discovery=False)


# ----------------------- Listing (metadata only) -----------------------


def _parse_modified_time(raw: str) -> datetime:
    # Drive returns RFC 3339 like "2026-04-25T12:34:56.789Z".
    return datetime.fromisoformat(raw.replace("Z", "+00:00"))


def _file_to_meta(f: dict) -> DriveDocMeta:
    return DriveDocMeta(
        file_id=f["id"],
        name=f["name"],
        mime_type=f["mimeType"],
        modified_time=_parse_modified_time(f["modifiedTime"]),
        parents=list(f.get("parents") or []),
    )


def list_team_docs(
    folder_id: str | None = None,
    *,
    service: Any | None = None,
    recursive: bool = True,
) -> list[DriveDocMeta]:
    """List supported docs under a folder. Cheap: metadata only, no body fetch."""
    folder_id = folder_id or _drive_secrets().get("shared_drive_folder_id")
    if not folder_id:
        raise RuntimeError("No folder_id provided or configured.")
    svc = service or _drive_service()
    files = svc.files()

    out: list[DriveDocMeta] = []
    folders_to_scan = [folder_id]
    seen_folders: set[str] = set()

    while folders_to_scan:
        f_id = folders_to_scan.pop()
        if f_id in seen_folders:
            continue
        seen_folders.add(f_id)

        page_token = None
        while True:
            resp = files.list(
                q=f"'{f_id}' in parents and trashed = false",
                fields=f"nextPageToken, files({_FILE_FIELDS})",
                pageSize=100,
                pageToken=page_token,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
            ).execute()

            for entry in resp.get("files", []):
                mt = entry["mimeType"]
                if mt == GOOGLE_FOLDER_MIME:
                    if recursive:
                        folders_to_scan.append(entry["id"])
                    continue
                if mt in SUPPORTED_MIMES:
                    out.append(_file_to_meta(entry))

            page_token = resp.get("nextPageToken")
            if not page_token:
                break

    return out


# ----------------------- Extractors (pure) -----------------------


def extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n\n".join(p.strip() for p in parts if p.strip())


def extract_docx_text(data: bytes) -> str:
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text(data: bytes, mime_type: str) -> str:
    """Dispatch to the right extractor. Google Doc bodies arrive as text/plain bytes."""
    if mime_type == GOOGLE_DOC_MIME or mime_type.startswith("text/"):
        return data.decode("utf-8", errors="replace").strip()
    if mime_type == PDF_MIME:
        return extract_pdf_text(data)
    if mime_type == DOCX_MIME:
        return extract_docx_text(data)
    raise ValueError(f"Unsupported mime type: {mime_type}")


# ----------------------- Download (body fetch) -----------------------


def _fetch_bytes(file_id: str, mime_type: str, *, service: Any) -> bytes:
    """Download a file's bytes. Google Docs are exported; binary types are read raw."""
    from googleapiclient.http import MediaIoBaseDownload

    files = service.files()
    if mime_type == GOOGLE_DOC_MIME:
        request = files.export_media(fileId=file_id, mimeType="text/plain")
    else:
        request = files.get_media(fileId=file_id, supportsAllDrives=True)

    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _status, done = downloader.next_chunk()
    return buf.getvalue()


def download_doc(file_id: str, *, service: Any | None = None) -> DriveDoc:
    """Fetch + extract a single doc by file_id."""
    svc = service or _drive_service()
    meta = svc.files().get(
        fileId=file_id,
        fields=_FILE_FIELDS,
        supportsAllDrives=True,
    ).execute()

    if meta["mimeType"] not in SUPPORTED_MIMES:
        raise ValueError(f"Unsupported mime type for {meta['name']}: {meta['mimeType']}")

    data = _fetch_bytes(file_id, meta["mimeType"], service=svc)
    text = extract_text(data, meta["mimeType"])

    return DriveDoc(
        file_id=meta["id"],
        name=meta["name"],
        mime_type=meta["mimeType"],
        modified_time=_parse_modified_time(meta["modifiedTime"]),
        parents=list(meta.get("parents") or []),
        text=text,
    )


# ----------------------- Change detection -----------------------


def changed_since(
    cutoff: datetime,
    *,
    folder_id: str | None = None,
    service: Any | None = None,
    recursive: bool = True,
) -> list[DriveDocMeta]:
    """Return docs whose modifiedTime is strictly after `cutoff`."""
    all_docs = list_team_docs(folder_id, service=service, recursive=recursive)
    return [d for d in all_docs if d.modified_time > cutoff]
